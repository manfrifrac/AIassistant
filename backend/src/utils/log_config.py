import logging
import json
from pprint import pformat
from typing import Any
from rich.console import Console
from rich.theme import Theme
from rich.logging import RichHandler
import os
from datetime import datetime
from rich.filesize import decimal
from rich.text import Text
from rich.console import ConsoleRenderable
from colorama import init
from docx import Document
from docx.shared import RGBColor

# Initialize colorama
init()

# Custom theme for rich
custom_theme = Theme({
    "info": "green",
    "warning": "yellow",
    "error": "red",
    "debug": "cyan"
})

console = Console(theme=custom_theme)

class PrettyFormatter(logging.Formatter):
    """Formatter che formatta correttamente dizionari, liste e altri oggetti complessi."""

    def _format_dict(self, d: dict, indent: int = 2, initial_indent: int = 0) -> str:
        """Format dictionary with proper indentation and line breaks."""
        try:
            formatted = json.dumps(d, indent=indent, ensure_ascii=False)
            lines = formatted.split('\n')
            
            # Calculate base indentation
            base_indent = ' ' * initial_indent
            
            # Format first line
            result = [lines[0]]
            
            # Format subsequent lines with increasing indentation
            if len(lines) > 1:
                # Add base indentation plus 2 spaces for nested items
                result.extend([base_indent + '  ' + line for line in lines[1:]])
            
            return '\n'.join(result)
        except Exception:
            return str(d)

    def _find_and_format_dicts(self, text: str) -> str:
        """Find and format dictionaries within text."""
        try:
            parts = []
            current_pos = 0
            
            # Find dictionary start positions
            start = text.find('{')
            while start != -1:
                # Add text before dictionary
                parts.append(text[current_pos:start])
                
                # Find matching closing brace
                count = 1
                end = start + 1
                while count > 0 and end < len(text):
                    if text[end] == '{':
                        count += 1
                    elif text[end] == '}':
                        count -= 1
                    end += 1
                
                if count == 0:  # Found complete dict
                    dict_str = text[start:end]
                    try:
                        # Try to parse and format the dictionary
                        d = eval(dict_str)
                        if isinstance(d, dict):
                            # Calculate initial indent based on text before dict
                            last_newline = text.rfind('\n', 0, start)
                            if (last_newline == -1):
                                initial_indent = start
                            else:
                                initial_indent = start - last_newline - 1
                            formatted_dict = self._format_dict(d, indent=4, initial_indent=initial_indent)
                            parts.append(formatted_dict)
                        else:
                            parts.append(dict_str)
                    except:
                        parts.append(dict_str)
                    
                    current_pos = end
                    start = text.find('{', current_pos)
                else:
                    # No matching brace found, add remaining text
                    parts.append(text[current_pos:])
                    break
            
            # Add any remaining text
            if current_pos < len(text):
                parts.append(text[current_pos:])
            
            return ''.join(parts)
        except Exception:
            return text

    def format(self, record: logging.LogRecord) -> str:
        try:
            # Get the basic message
            if isinstance(record.msg, dict):
                message = self._format_dict(record.msg)
            else:
                message = record.getMessage()
            
            # Format any dictionaries within the message
            message = self._find_and_format_dicts(message)
            
            return f"{record.levelname}: {message}"
        except Exception as e:
            return f"Error formatting log: {str(e)}"

class DocxFileHandler(logging.FileHandler):
    """Custom file handler that writes rich formatted text to a .docx file."""
    def __init__(self, filename, mode='a', encoding='utf-8'):
        super().__init__(filename, mode, encoding=encoding)
        self.document = Document()

    def emit(self, record):
        try:
            msg = self.format(record)
            self._write_to_docx(msg, record.levelname)
            self.flush()
        except Exception:
            self.handleError(record)

    def _write_to_docx(self, msg, levelname):
        """Write the log message to the .docx file with color formatting."""
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(msg)
        if levelname == "DEBUG":
            run.font.color.rgb = RGBColor(0, 255, 255)  # Cyan
        elif levelname == "INFO":
            run.font.color.rgb = RGBColor(0, 255, 0)  # Green
        elif levelname == "WARNING":
            run.font.color.rgb = RGBColor(255, 255, 0)  # Yellow
        elif levelname == "ERROR":
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        self.document.save(self.baseFilename)

def setup_logging(global_debug_mode: bool = False):
    """Setup logging configuration."""
    # Clear existing handlers first
    logging.getLogger().handlers = []
    
    # Set root logger level
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG if global_debug_mode else logging.INFO)

    # Ensure logs directory exists
    logs_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'logs')
    os.makedirs(logs_dir, exist_ok=True)

    # Create docx file handler with timestamp
    timestamp = datetime.now().strftime('%Y-%m-%d')
    log_file = os.path.join(logs_dir, f'app-log-{timestamp}.docx')
    file_handler = DocxFileHandler(log_file, mode='w', encoding='utf-8')
    file_handler.setLevel(logging.DEBUG if global_debug_mode else logging.INFO)
    file_handler.setFormatter(PrettyFormatter())

    # Console handler setup
    console_handler = RichHandler(
        console=console,
        show_time=False,
        show_level=True,
        show_path=True,
        rich_tracebacks=True,
        markup=False
    )
    console_handler.setLevel(logging.DEBUG if global_debug_mode else logging.INFO)
    console_handler.setFormatter(PrettyFormatter())

    # Add handlers to root logger
    root_logger.addHandler(file_handler)
    root_logger.addHandler(console_handler)

    # Specific logger configurations
    loggers_config = {
        "openai": logging.WARNING,
        "httpcore": logging.WARNING,
        "httpx": logging.WARNING,
        "gtts": logging.WARNING,
        "urllib3": logging.WARNING,
        "LangGraphSetup": logging.INFO,
        "StateManager": logging.DEBUG,
        "VoiceAssistant": logging.DEBUG,
        "SupervisorAgent": logging.DEBUG,
        "ResearcherAgent": logging.DEBUG,
        "GreetingAgent": logging.DEBUG,
        "LLMTools": logging.WARNING,
        "PythonREPLTool": logging.WARNING,
        "SpotifyTools": logging.WARNING,
        "TTS": logging.WARNING,
        "AudioHandler": logging.WARNING,
        "ErrorHandler": logging.WARNING,
        "botocore": logging.WARNING,
        "botocore.retryhandler": logging.WARNING,
        "pydub.converter": logging.WARNING,
        "watchfiles": logging.WARNING
    }

    # Configure specific loggers
    for logger_name, level in loggers_config.items():
        logger = logging.getLogger(logger_name)
        logger.setLevel(level)
        # Prevent duplicate logs by disabling propagation
        logger.propagate = False
        # Clear existing handlers
        logger.handlers = []
        # Add our handlers
        logger.addHandler(console_handler)
        logger.addHandler(file_handler)

    return root_logger

def get_logger(name: str) -> logging.Logger:
    """Helper per ottenere un logger con le impostazioni centralizzate."""
    return logging.getLogger(name)

# Remove global setup_logging call from here
# Let main.py handle the initialization
logger = get_logger(__name__)
